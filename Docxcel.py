from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from tkinter.ttk import Combobox
import docx
from openpyxl import load_workbook

root = Tk()
root.geometry("250x100")
root.resizable(0,0)
root.title("懶人程式")

def export_file():
    #print(excel_path.get())
    #print(word_path.get())
    #print(combo.get())
    #print(rad_selected.get())

    list_name = []  # 名稱
    list_respn = []  # 負責單位
    list_safeman = []  # 安全管理員
    list_room = []  # 所屬機房名稱
    list_month = []  # 月份
    list_filename = []  # 存檔檔名
    list_water = []  # 流水號

    try:
        wb = load_workbook(excel_path.get())
        ws = wb.active
    except:
        #print("Excel檔案未載入")
        mesg_excel_error()

    try:
        # 存入[]
        for i in range(2, ws.max_row+1):
            for j in range(1, 2):
                # print(ws.cell(row=i,column=j).value)
                list_name.append(ws.cell(row=i, column=j).value)
            for j in range(2, 3):
                list_respn.append(ws.cell(row=i, column=j).value)
            for j in range(3, 4):
                list_safeman.append(ws.cell(row=i, column=j).value)
            for j in range(4, 5):
                list_room.append(ws.cell(row=i, column=j).value)
            for j in range(5, 6):
                list_month.append(ws.cell(row=i, column=j).value)
            for j in range(6, 7):
                list_filename.append(ws.cell(row=i, column=j).value)
            for j in range(7, 8):
                list_water.append(ws.cell(row=i, column=j).value)
    except:
        #print("格式錯誤")
        mesg_format_error()

    try:
        doc = docx.Document(word_path.get())
        tables = doc.tables
        table = tables[0]
        i = 2

        for j in range(0,ws.max_row-1):
            table.cell(i, 0).text = list_name[j]
            table.cell(i, 1).text = str(list_respn[j])
            table.cell(i, 2).text = list_safeman[j]
            table.cell(i, 3).text = list_room[j]
            table.cell(i, 4).text = list_month[j]
            doc.save(folderPath.get()+'/' + list_filename[j] + '.docx')

        '''
        table.cell(i, 0).text = list_name[0]
        table.cell(i, 1).text = str(list_respn[0])
        table.cell(i, 2).text = list_safeman[0]
        table.cell(i, 3).text = list_room[0]
        table.cell(i, 4).text = list_month[0]
        doc.save('Test/'+list_filename[0]+'.docx')
        '''

    except:
        #print("Word尚未載入!")
        mesg_word_error()

def export_file2():

    list_name = []  # 名稱
    list_respn = []  # 負責單位
    list_safeman = []  # 安全管理員
    list_room = []  # 所屬機房名稱
    list_month = []  # 月份
    list_filename = []  # 存檔檔名
    list_water = []  # 流水號

    try:
        wb = load_workbook(excel_path.get())
        ws = wb.active
    except:
        mesg_excel_error()

    try:
        # 存入[]
        for i in range(2, ws.max_row+1):
            for j in range(1, 2):
                # print(ws.cell(row=i,column=j).value)
                list_name.append(ws.cell(row=i, column=j).value)
            for j in range(2, 3):
                list_respn.append(ws.cell(row=i, column=j).value)
            for j in range(3, 4):
                list_safeman.append(ws.cell(row=i, column=j).value)
            for j in range(4, 5):
                list_room.append(ws.cell(row=i, column=j).value)
            for j in range(5, 6):
                list_month.append(ws.cell(row=i, column=j).value)
            for j in range(6, 7):
                list_filename.append(ws.cell(row=i, column=j).value)
            for j in range(7, 8):
                list_water.append(ws.cell(row=i, column=j).value)
    except:
        #print("格式錯誤")
        mesg_format_error()

    if rad_selected.get() == 0:
        try:
            doc = docx.Document(word_path.get())

            for j in range(0,ws.max_row-1):
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph.text == "Test01":
                                    paragraph.text = list_name[j]
                                if paragraph.text == "Test02":
                                    paragraph.text = str(list_respn[j])
                                if paragraph.text == "Test03":
                                    paragraph.text = list_safeman[j]
                                if paragraph.text == "Test04":
                                    paragraph.text = list_room[j]
                                if paragraph.text == "Test05":
                                    paragraph.text = list_month[j]

                                # 還沒設定
                                if paragraph.text == "Test06":
                                    paragraph.text = "123"
                                if paragraph.text == "Test07":
                                    paragraph.text = "123"

                doc.save(folderPath.get()+'/' + list_filename[j] + '.docx')


                '''
                table.cell(i, 0).text = list_name[j]
                table.cell(i, 1).text = str(list_respn[j])
                table.cell(i, 2).text = list_safeman[j]
                table.cell(i, 3).text = list_room[j]
                table.cell(i, 4).text = list_month[j]
                doc.save(folderPath.get()+'/' + list_filename[j] + '.docx')
                '''
            mesg_done()
        except:
            #print("Word尚未載入!")
            mesg_word_error()


    elif rad_selected.get() == 1:
        wb02 = load_workbook(excel_ex_path.get())
        ws02 = wb02.active
        for j in range(0, ws.max_row - 1):

            for row in ws02.rows:
                for cell in row:
                    #print(cell.value, end=" ")
                    if cell.value == "Test01":
                        cell.value = list_name[j]
                    if cell.value == "Test02":
                        cell.value = str(list_respn[j])
                    if cell.value == "Test03":
                        cell.value = list_safeman[j]
                    if cell.value == "Test04":
                        cell.value = list_room[j]
                    if cell.value == "Test05":
                        cell.value = list_month[j]

                    # 還沒設定
                    if cell.value == "Test06":
                        cell.value = "123"
                    if cell.value == "Test07":
                        cell.value = "123"
            wb02.save(folderPath.get() + '/' + list_filename[j] + '.xlsx')
            #wb02.save('/Test/' + list_filename[j] + '.xlsx')
            #wb02.save("689.xlsx")
            mesg_done()

# Excel 資料檔
def open_excel():
    yxl_selected = filedialog.askopenfilename(initialdir="/", title="Select file",
                                               filetypes=(("xlsx files", "*.xlsx"), ("all files", "*.*")))
    excel_path.set(yxl_selected)

# Word 範本檔
def open_word():
    doc_selected = filedialog.askopenfilename(initialdir="/", title="Select file",
                                               filetypes=(("docx files", "*.docx"), ("all files", "*.*")))
    word_path.set(doc_selected)

# Excel 範本檔
def open_excel_ex():
    yxl_selected = filedialog.askopenfilename(initialdir="/", title="Select file",
                                               filetypes=(("xlsx files", "*.xlsx"), ("all files", "*.*")))
    excel_ex_path.set(yxl_selected)

def rad_edit_combobox():
    # Word
    if rad_selected.get()==0:
        combo['values']= ("機房安全管理作業日誌【安全維護區】(每週)","電信機房安全管理檢查表(每半年)")
    # Excel
    elif rad_selected.get()==1:
        combo['values'] = ("機房資安符合性檢查表(含安全維護區)(每年)","機房安全管理作業日誌【每月】","機房安全管理作業日誌【每季】","機房安全管理作業日誌【每半年】","機房安全管理作業日誌【每年】")



def get_dest_path():
    folder_selected = filedialog.askdirectory()
    folderPath.set(folder_selected)

def mesg_des():
    messagebox.showinfo("使用說明","請先載入檔案,才按下輸出!")

def mesg_excel_error():
    messagebox.showerror('錯誤', 'Excel尚未載入!')

def mesg_word_error():
    messagebox.showerror('錯誤', 'Word尚未載入!')

def mesg_format_error():
    messagebox.showerror('錯誤', '格式錯誤!')

def mesg_done():
    messagebox.showinfo("完成","輸出完畢!")


excel_path = StringVar()    # excel 資源檔路徑
word_path = StringVar()     # word 範本檔路徑
excel_ex_path = StringVar() # excel 範本檔路徑
folderPath = StringVar()    # 檔案輸出存放路徑
rad_selected = IntVar()     # 預設為Word

menu_bar = Menu(root)
file_menu = Menu(menu_bar,tearoff = 0)
des_menu = Menu(menu_bar,tearoff = 0)

menu_bar.add_cascade(label = "檔案",menu=file_menu)
file_menu.add_command(label="載入_Excel資源檔",command=open_excel)
file_menu.add_command(label="載入_Word範本",command=open_word)
file_menu.add_command(label="載入_Excel範本",command=open_excel_ex)

menu_bar.add_command(label="使用說明",command= mesg_des)

root.config(menu = menu_bar,bg='#81C0C0')


combo = Combobox(root,width=30)
combo['values'] = ("機房安全管理作業日誌【安全維護區】(每週)","電信機房安全管理檢查表(每半年)")

combo.current(0)
combo.place(x=5,y=5)

rad_ex_word = Radiobutton(root,text="Word範本",value=0,bg='#81C0C0',variable=rad_selected,command=rad_edit_combobox)
rad_ex_word.place(x=5,y=30)

rad_ex_excel = Radiobutton(root,text="Excel範本",value=1,bg='#81C0C0',variable=rad_selected,command=rad_edit_combobox)
rad_ex_excel.place(x=100,y=30)

btn_exp = Button(root, text="檔案輸出",command=export_file)
btn_exp.place(x=100,y=60)

btn_dest_path = Button(root, text="存放位置",command=get_dest_path)
btn_dest_path.place(x=5,y=60)

btn_exp2 = Button(root, text="檔案輸出2",command=export_file2)
btn_exp2.place(x=180,y=60)




'''
# 讀取 Excel內容
wb = load_workbook('達美樂.xlsx')
ws = wb.active

list_name = []  # 名稱
list_respn = [] # 負責單位
list_safeman = [] # 安全管理員
list_room = [] # 所屬機房名稱
list_month = [] # 月份
list_filename = [] # 存檔檔名
list_water = [] # 流水號

# 存入[]
for i in range(2,ws.max_row):
    for j in range(1,2):
        #print(ws.cell(row=i,column=j).value)
        list_name.append(ws.cell(row=i,column=j).value)
    for j in range(2,3):
        list_respn.append(ws.cell(row=i,column=j).value)
    for j in range(3, 4):
        list_safeman.append(ws.cell(row=i, column=j).value)
    for j in range(4, 5):
        list_room.append(ws.cell(row=i, column=j).value)
    for j in range(5, 6):
        list_month.append(ws.cell(row=i, column=j).value)
    for j in range(6, 7):
        list_filename.append(ws.cell(row=i, column=j).value)
    for j in range(7, 8):
        list_water.append(ws.cell(row=i, column=j).value)

doc = docx.Document('範本.docx')

tables = doc.tables
table = tables[0]

i = 2
#result = table.cell(i,0).text + " " + table.cell(i,1).text + " " + table.cell(i,2).text + " " + table.cell(i,3).text
table.cell(i, 0).text = list_name[0]
table.cell(i, 1).text = str(list_respn[0])
table.cell(i, 2).text = list_safeman[0]
table.cell(i, 3).text = list_room[0]
table.cell(i, 4).text = list_month[0]
 
doc.save('helloWorld.docx')
'''

root.mainloop()